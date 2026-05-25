from io import BytesIO
from pathlib import PurePath
import re

from docx import Document
from pypdf import PdfReader


TEXT_CONTENT_TYPES = {
    "text/plain",
    "text/markdown",
    "text/csv",
    "application/json",
}

ALLOWED_RESUME_EXTENSIONS = {".txt", ".md", ".csv", ".json", ".pdf", ".docx", ".png", ".jpg", ".jpeg", ".webp"}


class UnsupportedResumeFormatError(ValueError):
    """Raised when a resume format cannot be converted into text locally."""


def sanitize_resume_filename(filename: str | None) -> str:
    raw_name = PurePath((filename or "resume").replace("\\", "/")).name
    safe_name = re.sub(r"[^\w.\-\u4e00-\u9fff]+", "_", raw_name, flags=re.UNICODE).strip("._")
    if not safe_name:
        safe_name = "resume"
    return safe_name[:120]


def validate_resume_filename(filename: str | None) -> str:
    safe_name = sanitize_resume_filename(filename)
    extension = PurePath(safe_name).suffix.lower()
    if extension not in ALLOWED_RESUME_EXTENSIONS:
        raise UnsupportedResumeFormatError("unsupported_resume_format")
    return safe_name


def extract_resume_text(
    filename: str | None,
    content_type: str | None,
    payload: bytes,
    ocr_provider: str = "none",
) -> str:
    normalized_name = (filename or "").lower()
    normalized_type = (content_type or "").split(";")[0].strip().lower()

    if normalized_type in TEXT_CONTENT_TYPES or normalized_name.endswith((".txt", ".md", ".csv", ".json")):
        return _decode_text(payload)
    if normalized_type == "application/pdf" or normalized_name.endswith(".pdf"):
        return _extract_pdf_text(payload)
    if (
        normalized_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        or normalized_name.endswith(".docx")
    ):
        return _extract_docx_text(payload)
    if normalized_type.startswith("image/") or normalized_name.endswith((".png", ".jpg", ".jpeg", ".webp")):
        return _extract_image_text(payload, ocr_provider)
    raise UnsupportedResumeFormatError("unsupported_resume_format")


def normalize_extracted_text(text: str, max_chars: int = 12000) -> str:
    compacted = re.sub(r"\s+", " ", text).strip()
    return compacted[:max_chars]


def build_material_summary(
    *,
    interview_type: str,
    resume_text: str | None,
    job_title: str | None,
    job_requirements: str | None,
    target_school: str | None,
    major: str | None,
    research_direction: str | None,
) -> str:
    if interview_type == "job":
        resume_part = _preview(resume_text, 120)
        jd_part = _preview(job_requirements, 160)
        return f"目标岗位：{job_title or '未填写'}；岗位要求：{jd_part}；简历摘要：{resume_part}"
    direction = research_direction or "暂未填写具体方向"
    return f"目标院校：{target_school or '未填写'}；报考专业：{major or '未填写'}；研究方向：{direction}"


def extract_keywords(*texts: str | None, limit: int = 10) -> list[str]:
    joined = " ".join(text for text in texts if text)
    priority_keywords = [
        "FastAPI",
        "PostgreSQL",
        "Redis",
        "Docker",
        "Docker Compose",
        "模型路由",
        "日志追踪",
        "缓存",
        "接口稳定性",
        "OCR",
        "自然语言处理",
        "智能教育",
        "科研",
        "导师沟通",
    ]
    keywords: list[str] = []
    for keyword in priority_keywords:
        if keyword.lower() in joined.lower() and keyword not in keywords:
            keywords.append(keyword)

    for token in re.findall(r"[A-Za-z][A-Za-z0-9+#.-]{1,}|[\u4e00-\u9fff]{2,8}", joined):
        cleaned = token.strip("，。；：、,.!?;:()[]")
        if 2 <= len(cleaned) <= 24 and cleaned not in keywords:
            keywords.append(cleaned)
        if len(keywords) >= limit:
            break
    return keywords[:limit]


def _decode_text(payload: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return payload.decode(encoding)
        except UnicodeDecodeError:
            continue
    return payload.decode("utf-8", errors="ignore")


def _extract_pdf_text(payload: bytes) -> str:
    reader = PdfReader(BytesIO(payload))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_docx_text(payload: bytes) -> str:
    document = Document(BytesIO(payload))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    table_cells = [
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
        if cell.text.strip()
    ]
    return "\n".join([*paragraphs, *table_cells])


def _extract_image_text(payload: bytes, ocr_provider: str) -> str:
    provider = ocr_provider.strip().lower()
    if provider != "rapidocr":
        raise UnsupportedResumeFormatError("image_resume_ocr_not_configured")
    try:
        from rapidocr import RapidOCR  # type: ignore[import-not-found]
    except ImportError as exc:
        raise UnsupportedResumeFormatError("image_resume_ocr_dependency_missing") from exc

    result = RapidOCR()(payload)
    texts = _rapidocr_texts(result)
    if not texts:
        raise UnsupportedResumeFormatError("resume_text_empty")
    return "\n".join(texts)


def _rapidocr_texts(result: object) -> list[str]:
    if isinstance(result, tuple):
        result = result[0]
    txts = getattr(result, "txts", None)
    if txts:
        return [str(text) for text in txts if str(text).strip()]
    if not isinstance(result, list):
        return []
    texts: list[str] = []
    for item in result:
        if isinstance(item, (list, tuple)) and len(item) >= 2 and str(item[1]).strip():
            texts.append(str(item[1]))
    return texts


def _preview(text: str | None, max_chars: int) -> str:
    if not text:
        return "未提取到文本"
    return normalize_extracted_text(text, max_chars=max_chars)
